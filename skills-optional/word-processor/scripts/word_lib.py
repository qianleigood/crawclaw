#!/usr/bin/env python3
"""
Word 文档处理核心库
支持创建、读取、编辑、转换 Word 文档 (.docx)
"""

import os
import json
import re
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("错误：请安装依赖 - pip3 install python-docx")
    raise


class WordDocument:
    """Word 文档处理类"""
    
    def __init__(self, path: Optional[str] = None):
        """
        初始化文档
        
        Args:
            path: 文档路径，如果为 None 则创建新文档
        """
        if path and os.path.exists(path):
            self.doc = Document(path)
            self.path = path
        else:
            self.doc = Document()
            self.path = None
        
        # 设置默认中文字体
        self._set_default_font()
    
    def _set_default_font(self):
        """设置默认中文字体"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ========== 创建功能 ==========
    
    def add_heading(self, text: str, level: int = 1) -> None:
        """添加标题"""
        self.doc.add_heading(text, level)
    
    def add_paragraph(self, text: str, style: Optional[str] = None) -> None:
        """添加段落"""
        if style:
            self.doc.add_paragraph(text, style=style)
        else:
            para = self.doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(6)
    
    def add_table(self, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> None:
        """
        添加表格
        
        Args:
            rows: 行数
            cols: 列数
            data: 表格数据（可选）
        """
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        if data:
            for i, row_data in enumerate(data):
                if i < rows:
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            table.cell(i, j).text = str(cell_text)
    
    def add_picture(self, image_path: str, width: Optional[float] = None) -> None:
        """
        添加图片
        
        Args:
            image_path: 图片路径
            width: 宽度（英寸），默认 4 英寸
        """
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"图片不存在：{image_path}")
        
        if width:
            self.doc.add_picture(image_path, width=Inches(width))
        else:
            self.doc.add_picture(image_path, width=Inches(4))
    
    def add_page_break(self) -> None:
        """添加分页符"""
        self.doc.add_page_break()
    
    # ========== 读取功能 ==========
    
    def get_text(self) -> str:
        """获取所有文本"""
        return '\n'.join([para.text for para in self.doc.paragraphs])
    
    def get_tables(self) -> List[List[List[str]]]:
        """获取所有表格数据"""
        tables = []
        for table in self.doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables
    
    def get_images(self) -> List[str]:
        """获取图片信息（返回图片关系 ID 列表）"""
        images = []
        for rel in self.doc.part.rels.values():
            if "image" in rel.target_ref:
                images.append(rel.target_ref)
        return images
    
    def get_metadata(self) -> Dict[str, Any]:
        """获取文档元数据"""
        return {
            'paragraphs': len(self.doc.paragraphs),
            'tables': len(self.doc.tables),
            'sections': len(self.doc.sections),
            'path': self.path
        }
    
    # ========== 编辑功能 ==========
    
    def replace_text(self, old_text: str, new_text: str) -> int:
        """
        替换文本
        
        Args:
            old_text: 原文本
            new_text: 新文本
            
        Returns:
            替换次数
        """
        count = 0
        for para in self.doc.paragraphs:
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)
                count += 1
        
        # 同时替换表格中的文本
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)
                        count += 1
        
        return count
    
    def append_text(self, text: str) -> None:
        """在文档末尾添加文本"""
        self.doc.add_paragraph(text)
    
    def insert_table_at_end(self, data: List[List[str]]) -> None:
        """在文档末尾添加表格"""
        rows = len(data)
        cols = len(data[0]) if data else 0
        self.add_table(rows, cols, data)
    
    # ========== 模板填充 ==========
    
    def fill_template(self, data: Dict[str, Any]) -> int:
        """
        填充模板（替换 {{变量}} 占位符）
        
        Args:
            data: 数据字典
            
        Returns:
            替换次数
        """
        count = 0
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            count += self.replace_text(placeholder, str(value))
        return count
    
    def fill_template_from_json(self, json_path: str) -> int:
        """从 JSON 文件加载数据并填充模板"""
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return self.fill_template(data)
    
    # ========== 保存功能 ==========
    
    def save(self, path: Optional[str] = None) -> None:
        """
        保存文档
        
        Args:
            path: 保存路径，如果为 None 则使用原路径
        """
        save_path = path or self.path
        if not save_path:
            raise ValueError("必须指定保存路径")
        
        self.doc.save(save_path)
        self.path = save_path
    
    # ========== 静态方法 ==========
    
    @staticmethod
    def create_from_template(template_path: str, data: Dict[str, Any], output_path: str) -> 'WordDocument':
        """
        从模板创建文档
        
        Args:
            template_path: 模板路径
            data: 数据字典
            output_path: 输出路径
            
        Returns:
            WordDocument 实例
        """
        doc = WordDocument(template_path)
        doc.fill_template(data)
        doc.save(output_path)
        return doc
    
    @staticmethod
    def batch_create(template_path: str, data_list: List[Dict[str, Any]], output_dir: str) -> List[str]:
        """
        批量创建文档
        
        Args:
            template_path: 模板路径
            data_list: 数据列表
            output_dir: 输出目录
            
        Returns:
            生成的文件路径列表
        """
        os.makedirs(output_dir, exist_ok=True)
        output_files = []
        
        for i, data in enumerate(data_list):
            # 生成文件名（使用第一个字段的值或序号）
            filename = data.get('filename', f"document_{i+1}.docx")
            output_path = os.path.join(output_dir, filename)
            
            doc = WordDocument.create_from_template(template_path, data, output_path)
            output_files.append(output_path)
        
        return output_files


# ========== 便捷函数 ==========

def create_document(title: str, content: str, output_path: str) -> str:
    """快速创建文档"""
    doc = WordDocument()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    doc.save(output_path)
    return output_path


def read_document(path: str, output_format: str = 'text') -> Union[str, List, Dict]:
    """读取文档"""
    doc = WordDocument(path)
    
    if output_format == 'text':
        return doc.get_text()
    elif output_format == 'tables':
        return doc.get_tables()
    elif output_format == 'metadata':
        return doc.get_metadata()
    else:
        return doc.get_text()


def edit_document(input_path: str, output_path: str, **kwargs) -> str:
    """编辑文档"""
    doc = WordDocument(input_path)
    
    if 'replace' in kwargs:
        # 替换文本 "旧：新"
        for pair in kwargs['replace']:
            old, new = pair.split(':', 1)
            doc.replace_text(old, new)
    
    if 'append' in kwargs:
        doc.append_text(kwargs['append'])
    
    doc.save(output_path)
    return output_path


def fill_template(template_path: str, data: Dict[str, Any], output_path: str) -> str:
    """填充模板"""
    doc = WordDocument(template_path)
    doc.fill_template(data)
    doc.save(output_path)
    return output_path


def batch_fill_template(template_path: str, data_list: List[Dict[str, Any]], output_dir: str) -> List[str]:
    """批量填充模板"""
    return WordDocument.batch_create(template_path, data_list, output_dir)
