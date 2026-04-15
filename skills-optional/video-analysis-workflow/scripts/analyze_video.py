#!/usr/bin/env python3
"""
视频分析工作流 - V3 完整版（95%阈值 + 暂停恢复机制）

流程：
1. M2 Pro 硬件压缩原视频到 480x270
2. 压缩版 TransNetV2 分镜检测（快）
3. 用时间码在原版上切割（画质好）
4. 上传片段获取 CDN URL
5. Gemini 视频理解分析
6. 完整性检查（95%阈值）
7. 用户决策 / 暂停恢复
8. 生成 Word 报告
"""

import os
import sys
import json
import gc
import subprocess
import requests
import argparse
import logging
import time
import re
from pathlib import Path
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# 配置
SKILL_DIR = Path(__file__).resolve().parent.parent
SKILLS_DIR = SKILL_DIR.parent
TRANSNETV2_SCRIPT = SKILLS_DIR / "transnetv2-scene-detect/scripts/run_transnetv2_light.py"
PREPROCESS_SCRIPT = SKILLS_DIR / "transnetv2-scene-detect/scripts/preprocess_video.py"
TRANSNETV2_WEIGHTS = SKILLS_DIR / "transnetv2-scene-detect" / "assets/weights/transnetv2-pytorch-weights.pth"
VIDEO_UNDERSTAND_DIR = SKILLS_DIR / "video-understand"
VIDEO_UNDERSTAND_CONFIG = VIDEO_UNDERSTAND_DIR / "config.json"
VIDEO_UNDERSTAND_ANALYZER = VIDEO_UNDERSTAND_DIR / "video_analyzer.py"
XHS_AUTO_IMPORT_DIR = SKILLS_DIR / "xhs-auto-import"
XHS_URL_UTIL = XHS_AUTO_IMPORT_DIR / "utils/url.py"
DOUYIN_PLATFORM = XHS_AUTO_IMPORT_DIR / "platforms/douyin.py"
XIAOHONGSHU_PLATFORM = XHS_AUTO_IMPORT_DIR / "platforms/xiaohongshu.py"
XHS_PYTHON = XHS_AUTO_IMPORT_DIR / ".venv/bin/python"
FEISHU_TOOLKIT_ENV = SKILLS_DIR / "feishu-office-toolkit/server/.env"


def get_workflow_missing_dependencies(input_value: str | None = None, feishu_send_enabled: bool = False) -> list[str]:
    """返回 workflow 当前缺失的本地依赖/资源。"""
    missing = []
    required_paths = [
        (TRANSNETV2_SCRIPT, "TransNetV2 检测脚本"),
        (PREPROCESS_SCRIPT, "TransNetV2 预处理脚本"),
        (TRANSNETV2_WEIGHTS, "TransNetV2 权重文件"),
        (VIDEO_UNDERSTAND_CONFIG, "video-understand 配置文件"),
        (VIDEO_UNDERSTAND_ANALYZER, "video-understand 分析脚本"),
    ]

    url = extract_supported_url(input_value) if input_value else None
    if url:
        required_paths.extend([
            (XHS_URL_UTIL, "xhs-auto-import URL 识别模块"),
        ])
        if "douyin" in url:
            required_paths.append((DOUYIN_PLATFORM, "xhs-auto-import 抖音平台模块"))
        if "xiaohongshu" in url or "xhslink" in url:
            required_paths.extend([
                (XIAOHONGSHU_PLATFORM, "xhs-auto-import 小红书平台模块"),
                (XHS_PYTHON, "xhs-auto-import 虚拟环境 Python"),
            ])

    if feishu_send_enabled:
        required_paths.append((FEISHU_TOOLKIT_ENV, "飞书工具包环境文件"))

    for path, label in required_paths:
        if not path.exists():
            missing.append(f"{label}: {path}")
    return missing


def ensure_workflow_dependencies(input_value: str | None = None, feishu_send_enabled: bool = False) -> None:
    missing = get_workflow_missing_dependencies(input_value=input_value, feishu_send_enabled=feishu_send_enabled)
    if missing:
        raise FileNotFoundError("workflow dependency preflight failed:\n- " + "\n- ".join(missing))


def get_video_understand_missing_envs() -> list[str]:
    """从 video-understand/config.json 中提取 `${ENV}` 占位，并返回缺失项。"""
    if not VIDEO_UNDERSTAND_CONFIG.exists():
        return []

    try:
        with open(VIDEO_UNDERSTAND_CONFIG, 'r', encoding='utf-8') as f:
            cfg = json.load(f)
    except Exception as e:
        print(f"⚠️  读取 video-understand 配置失败: {e}")
        return []

    placeholders = []
    pattern = re.compile(r'^\$\{([A-Z0-9_]+)\}$')
    for value in cfg.values():
        if isinstance(value, str):
            m = pattern.match(value)
            if m:
                placeholders.append(m.group(1))

    return [name for name in placeholders if not os.environ.get(name)]


def bootstrap_video_understand_env() -> None:
    missing = get_video_understand_missing_envs()
    if missing:
        print(f"⚠️  video-understand 依赖的环境变量未设置: {', '.join(missing)}")


bootstrap_video_understand_env()

# 导入过渡镜头检测器
sys.path.insert(0, str(Path(__file__).parent))
from transition_detector import TransitionDetector

# 成功率阈值 - 95%
SUCCESS_RATE_THRESHOLD = 0.95

# ============================================================================
# 日志配置
# ============================================================================

def setup_logging(output_dir: Path, video_id: str) -> logging.Logger:
    """设置日志系统"""
    logger = logging.getLogger(f"video_analyzer_{video_id}")
    logger.setLevel(logging.DEBUG)
    logger.handlers = []
    
    log_dir = output_dir / "logs"
    log_dir.mkdir(exist_ok=True)
    
    log_file = log_dir / f"analyze_{video_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    file_handler = logging.FileHandler(log_file, encoding='utf-8')
    file_handler.setLevel(logging.DEBUG)
    file_format = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s', datefmt='%Y-%m-%d %H:%M:%S')
    file_handler.setFormatter(file_format)
    logger.addHandler(file_handler)
    
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.INFO)
    console_format = logging.Formatter('%(message)s')
    console_handler.setFormatter(console_format)
    logger.addHandler(console_handler)
    
    logger.info(f"📝 日志文件: {log_file}")
    return logger


logger = logging.getLogger(__name__)


def get_progress_file(output_dir: Path) -> Path:
    """获取进度文件路径"""
    return output_dir / ".analysis_progress.json"


def get_incomplete_state_file(output_dir: Path) -> Path:
    """获取不完整状态文件路径"""
    return output_dir / ".incomplete_state.json"


# 导入视频理解模块
sys.path.insert(0, str(VIDEO_UNDERSTAND_DIR))
from video_analyzer import upload_file, analyze_video


# ============================================================================
# 进度管理（断点续传）
# ============================================================================

def load_progress(output_dir: Path) -> dict:
    """加载进度"""
    progress_file = get_progress_file(output_dir)
    if progress_file.exists():
        try:
            with open(progress_file, 'r') as f:
                return json.load(f)
        except:
            pass
    return {
        'video_path': '',
        'preprocessed': False,
        'scenes_detected': False,
        'scenes': [],
        'scenes_cut': 0,
        'analyses': [],
        'completed': False
    }


def save_progress(output_dir: Path, progress: dict):
    """保存进度"""
    progress_file = get_progress_file(output_dir)
    progress['timestamp'] = datetime.now().isoformat()
    with open(progress_file, 'w') as f:
        json.dump(progress, f, indent=2)


def clear_progress(output_dir: Path):
    """清除进度"""
    progress_file = get_progress_file(output_dir)
    if progress_file.exists():
        progress_file.unlink()


# ============================================================================
# 不完整状态管理（暂停恢复机制）
# ============================================================================

def save_incomplete_state(output_dir: Path, video_path: str, scenes: list, 
                          analyses: list, completeness: dict, user_prompt: str = None):
    """
    保存不完整状态，供下次恢复
    """
    state_file = get_incomplete_state_file(output_dir)
    
    state = {
        "video_path": video_path,
        "video_id": get_video_id(video_path),
        "scenes": scenes,
        "analyses": analyses,
        "completeness": completeness,
        "user_prompt": user_prompt,
        "paused_at": datetime.now().isoformat(),
        "message": f"分析暂停：成功率 {completeness['success_rate']*100:.1f}% (阈值 {SUCCESS_RATE_THRESHOLD*100}%)",
        "status": "paused"
    }
    
    with open(state_file, 'w', encoding='utf-8') as f:
        json.dump(state, f, ensure_ascii=False, indent=2)
    
    logger.info(f"💾 暂停状态已保存: {state_file}")
    logger.info(f"下次恢复: python3 scripts/analyze_video.py --resume")


def load_incomplete_state(output_dir: Path) -> dict:
    """加载不完整状态"""
    state_file = get_incomplete_state_file(output_dir)
    if state_file.exists():
        try:
            with open(state_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            logger.warning(f"加载暂停状态失败: {e}")
    return None


def clear_incomplete_state(output_dir: Path):
    """清除不完整状态"""
    state_file = get_incomplete_state_file(output_dir)
    if state_file.exists():
        state_file.unlink()
        logger.info("🗑️ 暂停状态已清除")


# ============================================================================
# 完整性检查
# ============================================================================

def check_analysis_completeness(analyses: list) -> dict:
    """
    检查分析完整性
    
    Returns:
        {
            "complete": bool,           # 是否完整（≥95%）
            "success_rate": float,      # 成功率
            "total": int,
            "success": int,
            "failed": int,
            "skipped": int,
            "can_retry": list           # 可重试的失败项
        }
    """
    total = len(analyses)
    success = len([a for a in analyses if a.get("status") == "success"])
    skipped = len([a for a in analyses if a.get("status") == "skipped"])
    failed = total - success - skipped
    success_rate = success / total if total > 0 else 0
    
    # 可重试的失败项（排除文件不存在、过大等）
    can_retry = [
        a for a in analyses 
        if a.get("status") in ["failed", "upload_failed", "exception", "error"]
        and Path(a.get("path", "")).exists()
    ]
    
    return {
        "complete": success_rate >= SUCCESS_RATE_THRESHOLD,
        "success_rate": success_rate,
        "total": total,
        "success": success,
        "failed": failed,
        "skipped": skipped,
        "can_retry": can_retry
    }


# ============================================================================
# 用户交互
# ============================================================================

def prompt_user_decision(completeness: dict, output_dir: Path) -> str:
    """
    提示用户决策（命令行交互）
    
    Returns:
        "continue" | "retry" | "pause" | "abort"
    """
    print("\n" + "="*60)
    print("⚠️  分析未达到 95% 完整性要求")
    print("="*60)
    print(f"总片段: {completeness['total']}")
    print(f"成功: {completeness['success']} ({completeness['success_rate']*100:.1f}%)")
    print(f"失败: {completeness['failed']}")
    print(f"跳过: {completeness['skipped']} (过大/不存在)")
    print(f"可重试: {len(completeness['can_retry'])}")
    print(f"\n阈值要求: {SUCCESS_RATE_THRESHOLD*100}%")
    print("="*60)
    
    print("\n请选择：")
    print("  [c] continue - 继续生成报告（含失败标记）")
    print("  [r] retry - 立即重试失败片段")
    print("  [p] pause - 暂停，下次启动继续分析")
    print("  [a] abort - 放弃，删除进度")
    
    while True:
        try:
            choice = input("\n你的选择 [c/r/p/a]: ").strip().lower()
            if choice in ['c', 'continue']:
                return "continue"
            elif choice in ['r', 'retry']:
                return "retry"
            elif choice in ['p', 'pause']:
                return "pause"
            elif choice in ['a', 'abort']:
                return "abort"
            else:
                print("无效选择，请重新输入")
        except (EOFError, KeyboardInterrupt):
            # 非交互环境，默认暂停
            print("\n[非交互环境，自动选择暂停]")
            return "pause"


def show_resume_prompt(state: dict) -> str:
    """
    显示恢复提示
    
    Returns:
        "continue" | "retry" | "new"
    """
    print("\n" + "="*60)
    print("📋 发现未完成的分析")
    print("="*60)
    print(f"视频: {state['video_path']}")
    print(f"暂停时间: {state['paused_at']}")
    print(f"状态: {state['message']}")
    
    comp = state['completeness']
    print(f"\n上次分析结果：")
    print(f"  成功率: {comp['success_rate']*100:.1f}%")
    print(f"  成功: {comp['success']}/{comp['total']}")
    print(f"  失败: {comp['failed']}")
    print(f"  可重试: {len(comp['can_retry'])}")
    
    print("\n请选择：")
    print("  [c] continue - 继续生成报告")
    print("  [r] retry - 重试失败片段")
    print("  [n] new - 开始新分析（忽略此状态）")
    
    while True:
        try:
            choice = input("\n你的选择 [c/r/n]: ").strip().lower()
            if choice in ['c', 'continue']:
                return "continue"
            elif choice in ['r', 'retry']:
                return "retry"
            elif choice in ['n', 'new']:
                return "new"
            else:
                print("无效选择，请重新输入")
        except (EOFError, KeyboardInterrupt):
            print("\n[非交互环境，自动选择继续生成报告]")
            return "continue"


# ============================================================================
# 视频处理函数
# ============================================================================

def extract_supported_url(text: str) -> str | None:
    patterns = [
        r'https?://(v\.douyin|www\.douyin|www\.iesdouyin)\.(com|cn)/\S+',
        r'https?://(www\.)?(xiaohongshu|xhslink)\.(com|cn)/\S+',
    ]
    for pattern in patterns:
        m = re.search(pattern, text)
        if m:
            return m.group(0)
    return None


def resolve_xiaohongshu_via_subprocess(url: str, download_dir: Path, xhs_skill_dir: Path) -> str:
    """通过 xhs-auto-import 自己的虚拟环境解析/下载小红书，避免依赖污染当前 venv。"""
    xhs_python = xhs_skill_dir / ".venv" / "bin" / "python"
    if not xhs_python.exists():
        raise FileNotFoundError(f"未找到 xhs-auto-import 虚拟环境: {xhs_python}")

    helper = r'''
import asyncio, json, shutil, sys
from pathlib import Path
sys.path.insert(0, sys.argv[1])
from platforms.xiaohongshu import XiaohongshuPlatform
url = sys.argv[2]
out = Path(sys.argv[3])
skill_dir = Path(sys.argv[1])
async def main():
    d = XiaohongshuPlatform()
    r = await d.fetch(url)
    internal_dir = skill_dir / '.hooks_output' / 'media' / 'xhs'
    r = await d.download_media(r, internal_dir)
    video_path = None
    if r.video_path and Path(r.video_path).exists():
        src = Path(r.video_path)
        out.mkdir(parents=True, exist_ok=True)
        dst = out / src.name
        if src.resolve() != dst.resolve():
            shutil.copy2(src, dst)
        video_path = str(dst)
    print(json.dumps({"video_path": video_path, "image_paths": r.image_paths}, ensure_ascii=False))
asyncio.run(main())
'''
    proc = subprocess.run(
        [str(xhs_python), "-c", helper, str(xhs_skill_dir), url, str(download_dir)],
        capture_output=True,
        text=True,
        timeout=180,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr.strip() or proc.stdout.strip() or "xiaohongshu helper failed")

    payload = None
    for line in reversed((proc.stdout or "").splitlines()):
        line = line.strip()
        if line.startswith("{") and line.endswith("}"):
            payload = json.loads(line)
            break
    if not payload:
        raise RuntimeError(f"xiaohongshu helper 无有效输出: {proc.stdout.strip()[:500]}")

    video_path = payload.get("video_path")
    if video_path and Path(video_path).exists():
        return video_path
    raise FileNotFoundError(f"小红书链接已解析，但未下载到本地视频: {url}")


def resolve_input_video(input_value: str, output_dir: Path) -> str:
    """将输入解析为本地视频文件路径。支持本地文件和分享文案/URL。"""
    input_path = Path(input_value)
    if input_path.exists():
        return str(input_path.resolve())

    url = extract_supported_url(input_value)
    if not url:
        return input_value

    download_dir = output_dir / "source_videos"
    download_dir.mkdir(parents=True, exist_ok=True)

    xhs_skill_dir = SKILLS_DIR / "xhs-auto-import"
    sys.path.insert(0, str(xhs_skill_dir))
    from utils.url import detect_platform
    platform = detect_platform(url)
    if not platform:
        raise ValueError(f"不支持的链接输入: {url}")

    if platform == 'xiaohongshu':
        return resolve_xiaohongshu_via_subprocess(url, download_dir, xhs_skill_dir)

    import asyncio
    if platform == 'douyin':
        from platforms.douyin import DouyinPlatform
        downloader = DouyinPlatform()
    else:
        raise ValueError(f"暂不支持的平台: {platform}")

    result = asyncio.run(downloader.fetch(url))
    result = asyncio.run(downloader.download_media(result, download_dir))
    if result.video_path and Path(result.video_path).exists():
        return result.video_path
    raise FileNotFoundError(f"链接已识别，但未下载到本地视频: {url}")


def get_video_id(video_path: str) -> str:
    """从视频路径或 URL 提取视频ID"""
    url = extract_supported_url(video_path)
    if url:
        m = re.search(r'(?:modal_id=|explore/|video/|note/|/)(\d{8,})', url)
        if m:
            return m.group(1)
    path = Path(video_path)
    stem = path.stem
    if stem.isdigit():
        return stem
    return stem


def preprocess_video(input_path: str, output_dir: Path, video_id: str) -> str:
    """步骤0: M2 Pro 硬件压缩视频到 480x270"""
    logger.info("="*60)
    logger.info("步骤0: M2 Pro 硬件压缩 (480x270)")
    logger.info("="*60)
    
    output_path = output_dir / f"preprocessed_{video_id}_480x270.mp4"
    
    if output_path.exists():
        logger.info(f"♻️  使用已有压缩版: {output_path}")
        return str(output_path)
    
    logger.info(f"📦 开始压缩视频: {input_path}")
    
    try:
        start_time = datetime.now()
        result = subprocess.run(
            [sys.executable, str(PREPROCESS_SCRIPT), input_path, "-o", str(output_path)],
            capture_output=True, text=True, timeout=120
        )
        elapsed = (datetime.now() - start_time).total_seconds()
        
        if output_path.exists():
            output_size = output_path.stat().st_size / 1024 / 1024
            logger.info(f"✅ 压缩完成: {output_path} ({output_size:.2f} MB, 耗时 {elapsed:.1f}s)")
            return str(output_path)
        else:
            logger.error(f"❌ 压缩失败，输出文件不存在")
            return input_path
    except subprocess.TimeoutExpired:
        logger.error(f"❌ 压缩超时（120秒）")
        return input_path
    except Exception as e:
        logger.error(f"❌ 压缩异常: {e}")
        return input_path


def detect_scenes(video_path: str, output_dir: Path, video_id: str = None) -> list:
    """步骤1: TransNetV2 分镜检测"""
    logger.info("\n" + "="*60)
    logger.info("步骤1: TransNetV2 分镜检测")
    logger.info("="*60)
    
    if video_id is None:
        video_id = get_video_id(video_path)
    scenes_file = output_dir / f"scenes_{video_id}.json"
    
    logger.info(f"🎬 视频ID: {video_id}")
    logger.info(f"📁 场景文件: {scenes_file}")
    
    result = subprocess.run(
        [sys.executable, str(TRANSNETV2_SCRIPT), video_path, "--weights", str(TRANSNETV2_WEIGHTS), "--output", str(scenes_file)],
        capture_output=True, text=True, timeout=300
    )
    
    if result.stdout:
        print(result.stdout)
    if result.stderr:
        print(result.stderr)
    
    if scenes_file.exists():
        with open(scenes_file) as f:
            data = json.load(f)
            scenes = data.get("scenes", data)
        logger.info(f"\n✅ 检测到 {len(scenes)} 个镜头")
        
        # 步骤1.5: 过渡镜头过滤（可选，默认关闭）
        # 注意：过渡镜头过滤可能会误过滤有效镜头，建议手动审核
        ENABLE_TRANSITION_FILTER = False  # 设置为True启用过滤
        
        if ENABLE_TRANSITION_FILTER:
            logger.info("\n" + "="*60)
            logger.info("步骤1.5: 过渡镜头过滤")
            logger.info("="*60)
            
            # 使用原始视频路径（非压缩版）进行更准确的分析
            original_video = output_dir.parent / f"{video_id}_original.mp4"
            if not original_video.exists():
                original_video = video_path
            
            detector = TransitionDetector()
            content_scenes, transition_scenes = detector.filter_scenes(
                str(original_video), scenes
            )
            
            logger.info(f"\n✅ 过滤后: {len(content_scenes)} 个有效镜头")
            
            # 保存过滤结果
            filtered_file = output_dir / f"scenes_{video_id}_filtered.json"
            with open(filtered_file, 'w') as f:
                json.dump({
                    'original_count': len(scenes),
                    'filtered_count': len(content_scenes),
                    'transition_count': len(transition_scenes),
                    'scenes': content_scenes,
                    'transition_scenes': transition_scenes
                }, f, indent=2)
            
            return content_scenes
        else:
            logger.info("\n⏩ 跳过过渡镜头过滤（已禁用）")
            return scenes
    
    logger.error("❌ 分镜检测失败，未生成场景文件")
    return []


def check_audio(video_path: Path) -> bool:
    """检查视频是否有音频流"""
    try:
        result = subprocess.run(
            ["ffprobe", "-v", "error", "-show_entries", "stream=codec_type",
             "-of", "default=noprint_wrappers=1", str(video_path)],
            capture_output=True, text=True, timeout=5
        )
        return "codec_type=audio" in result.stdout
    except:
        return False


def cut_scenes(original_video: str, scenes: list, output_dir: Path, video_id: str = None) -> list:
    """步骤2: 在原版视频上切割片段"""
    logger.info("\n" + "="*60)
    logger.info("步骤2: 在原版视频上切割片段")
    logger.info("="*60)
    
    if video_id is None:
        video_id = get_video_id(original_video)
    clips_dir = output_dir / f"clips_{video_id}"
    clips_dir.mkdir(exist_ok=True)
    
    logger.info(f"📁 片段目录: {clips_dir}")
    logger.info(f"🎯 需要切割 {len(scenes)} 个镜头")
    
    cut_scenes = []
    existing_count = 0
    new_count = 0
    failed_count = 0
    
    for i, scene in enumerate(scenes):
        output_path = clips_dir / f"scene_{i:03d}_{scene.get('start', 0):.3f}s.mp4"
        
        if output_path.exists():
            has_audio = check_audio(output_path)
            size_mb = output_path.stat().st_size / 1024 / 1024
            cut_scenes.append({
                "idx": i,
                "path": str(output_path),
                "scene": scene,
                "has_audio": has_audio,
                "size_mb": size_mb
            })
            existing_count += 1
            if (i + 1) % 50 == 0 or i == len(scenes) - 1:
                logger.info(f"   进度: {i+1}/{len(scenes)} (已存在: {existing_count}, 新建: {new_count}, 失败: {failed_count})")
            continue
        
        start = scene.get("start", scene.get("start_time", 0))
        end = scene.get("end", scene.get("end_time", start + 5))
        duration = min(end - start, 10)
        
        # 修复：使用重新编码而不是copy，避免关键帧导致的黑屏问题
        cmd = [
            "ffmpeg", "-y",
            "-ss", str(start),
            "-i", original_video,
            "-t", str(duration),
            "-map", "0:v:0",
            "-map", "0:a:0?",
            "-c:v", "libx264",
            "-crf", "23",
            "-preset", "fast",
            "-pix_fmt", "yuv420p",
            "-c:a", "aac",
            "-b:a", "128k",
            "-r", "30",
            "-avoid_negative_ts", "make_zero",
            str(output_path)
        ]
        
        try:
            result = subprocess.run(cmd, capture_output=True, timeout=60)
            
            if output_path.exists() and output_path.stat().st_size > 1000:
                size_mb = output_path.stat().st_size / 1024 / 1024
                has_audio = check_audio(output_path)
                cut_scenes.append({
                    "idx": i,
                    "path": str(output_path),
                    "scene": scene,
                    "has_audio": has_audio,
                    "size_mb": size_mb
                })
                new_count += 1
            else:
                failed_count += 1
        except Exception as e:
            failed_count += 1
            logger.debug(f"   片段 {i+1} 切割失败: {e}")
    
    logger.info(f"\n✅ 成功切割 {len(cut_scenes)} 个片段")
    return cut_scenes


# ============================================================================
# 上传和分析（带重试）
# ============================================================================

@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=4, max=10),
    retry=retry_if_exception_type((requests.exceptions.RequestException, Exception)),
    before_sleep=lambda retry_state: print(f"   上传失败，{retry_state.next_action.sleep}秒后重试...")
)
def upload_with_retry(scene_path: str) -> str:
    """带重试的上传"""
    return upload_file(scene_path, use_cache=True)


@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=4, max=10),
    retry=retry_if_exception_type((requests.exceptions.RequestException, Exception)),
    before_sleep=lambda retry_state: print(f"   分析失败，{retry_state.next_action.sleep}秒后重试...")
)
def analyze_with_retry(video_url: str, question: str) -> str:
    """带重试的分析"""
    return analyze_video(video_url, question)


def analyze_scene(scene_path: str, idx: int, total: int, scene: dict, question: str = None) -> dict:
    """分析单个场景"""
    print(f"   分析片段 {idx+1}/{total}...", end=' ', flush=True)
    
    if not os.path.exists(scene_path):
        print(f"❌ 文件不存在")
        return {
            "idx": idx,
            "path": scene_path,
            "scene": scene,
            "analysis": "[文件不存在]",
            "status": "error",
            "size_mb": 0
        }
    
    size_mb = os.path.getsize(scene_path) / 1024 / 1024
    if size_mb > 15:
        print(f"太大({size_mb:.1f}MB > 15MB)，跳过")
        return {
            "idx": idx,
            "path": scene_path,
            "scene": scene,
            "analysis": f"[片段过大 {size_mb:.1f}MB，跳过分析]",
            "status": "skipped",
            "size_mb": size_mb
        }
    
    # 上传
    try:
        video_url = upload_with_retry(scene_path)
    except Exception as e:
        print(f"❌ 上传失败: {str(e)[:50]}")
        return {
            "idx": idx,
            "path": scene_path,
            "scene": scene,
            "analysis": f"[上传失败: {str(e)[:100]}]",
            "status": "upload_failed",
            "size_mb": size_mb
        }
    
    if not video_url:
        print(f"❌ 上传返回空URL")
        return {
            "idx": idx,
            "path": scene_path,
            "scene": scene,
            "analysis": "[上传失败: 返回空URL]",
            "status": "upload_failed",
            "size_mb": size_mb
        }
    
    # 分析
    if question is None:
        question = """请分析这个视频镜头，并按以下结构输出：\n\n1. 镜头内容：这一镜头拍到了什么\n2. 景别与机位：远景/中景/近景，平视/俯拍/仰拍等\n3. 运镜方式：固定、推、拉、摇、移、跟拍等；如无法判断请说明\n4. 人物与动作：人物状态、动作、互动\n5. 场景与道具：环境、物件、视觉元素\n6. 字幕与声音：字幕文案、旁白、音乐、环境声\n7. 氛围与作用：这个镜头传达的情绪，以及在整体内容中的作用\n\n要求：\n- 用中文输出\n- 不要编造细节\n- 看不清就明确说明\n- 保持简洁、客观、可用于后续汇总成分镜脚本"""
    
    try:
        result = analyze_with_retry(video_url, question)
        if result and len(result) > 10:
            print("✅")
            return {
                "idx": idx,
                "path": scene_path,
                "scene": scene,
                "analysis": result,
                "status": "success",
                "size_mb": size_mb
            }
        else:
            print(f"⚠️ 结果太短({len(result) if result else 0}字符)")
            return {
                "idx": idx,
                "path": scene_path,
                "scene": scene,
                "analysis": "[分析结果异常: 内容太短]",
                "status": "failed",
                "size_mb": size_mb
            }
    except Exception as e:
        print(f"❌ 分析失败: {str(e)[:50]}")
        return {
            "idx": idx,
            "path": scene_path,
            "scene": scene,
            "analysis": f"[分析失败: {str(e)[:200]}]",
            "status": "failed",
            "size_mb": size_mb
        }


def analyze_all_scenes(cut_scenes: list, question: str = None, 
                       max_concurrent: int = 2, progress: dict = None,
                       output_dir: Path = None) -> list:
    """分析所有场景"""
    print("\n" + "="*60)
    print("步骤3: 上传并分析所有场景")
    print(f"   并发数：{max_concurrent}")
    print("="*60)
    
    analyses = progress.get('analyses', []) if progress else []
    existing_analyses = {a['idx']: a for a in analyses}
    
    # 过滤已分析的
    scenes_to_analyze = [s for s in cut_scenes if s['idx'] not in existing_analyses]
    
    if scenes_to_analyze:
        print(f"📦 共 {len(scenes_to_analyze)} 个新片段需要分析")
        
        with ThreadPoolExecutor(max_workers=max_concurrent) as executor:
            future_to_item = {
                executor.submit(analyze_scene, item["path"], item["idx"], 
                              len(cut_scenes), item["scene"], question): item
                for item in scenes_to_analyze
            }
            
            completed = 0
            for future in as_completed(future_to_item):
                item = future_to_item[future]
                try:
                    result = future.result()
                    existing_analyses[result['idx']] = result
                except Exception as e:
                    print(f"\n   片段 {item['idx']+1} 异常: {e}")
                    existing_analyses[item['idx']] = {
                        "idx": item['idx'],
                        "path": item['path'],
                        "scene": item['scene'],
                        "analysis": f"[分析异常: {e}]",
                        "status": "exception",
                        "size_mb": item.get('size_mb', 0)
                    }
                
                completed += 1
                if completed % 5 == 0 or completed == len(scenes_to_analyze):
                    success = len([a for a in existing_analyses.values() if a['status'] == 'success'])
                    print(f"   进度: {completed}/{len(scenes_to_analyze)} (成功:{success})")
                
                # 每完成一个片段就实时保存进度（修复断点续传问题）
                if progress is not None and output_dir is not None:
                    progress['analyses'] = list(existing_analyses.values())
                    save_progress(output_dir, progress)
                
                gc.collect()
    else:
        print(f"⏩ 所有 {len(cut_scenes)} 个片段已分析")
    
    # 返回完整列表
    return list(existing_analyses.values())


# ============================================================================
# 重试失败片段
# ============================================================================

def retry_failed_scenes(analyses: list, question: str, max_concurrent: int = 2) -> list:
    """重试失败的片段"""
    failed_items = [a for a in analyses if a['status'] in ['failed', 'upload_failed', 'exception', 'error']]
    
    if not failed_items:
        print("✅ 没有失败的片段需要重试")
        return analyses
    
    print(f"\n🔄 重试 {len(failed_items)} 个失败的片段...")
    
    # 更新状态
    for item in failed_items:
        print(f"   重试片段 {item['idx']+1}...", end=' ', flush=True)
        new_result = analyze_scene(item['path'], item['idx'], len(analyses), item['scene'], question)
        
        # 更新结果
        for i, a in enumerate(analyses):
            if a['idx'] == new_result['idx']:
                analyses[i] = new_result
                break
    
    print(f"\n✅ 重试完成")
    return analyses


# ============================================================================
# 分批分析（独立进程，避免超时）
# ============================================================================

def analyze_all_scenes_batched(cut_scenes: list, question: str, progress: dict, output_dir: Path) -> list:
    """
    分批分析所有场景，每批使用独立进程
    避免单进程运行时间过长被系统终止
    """
    print("\n" + "="*60)
    print("步骤3: 上传并分析所有场景（分批独立进程）")
    print("="*60)

    missing_envs = get_video_understand_missing_envs()
    if missing_envs:
        print(f"⚠️  video-understand 配置引用了环境变量，但当前未设置: {', '.join(missing_envs)}")
    
    analyses = progress.get('analyses', [])
    existing_analyses = {a['idx']: a for a in analyses}

    retryable_statuses = {'failed', 'upload_failed', 'exception', 'error'}
    pending_scenes = []
    for scene in cut_scenes:
        existing = existing_analyses.get(scene['idx'])
        if existing is None:
            pending_scenes.append(scene)
            continue
        if existing.get('status') in retryable_statuses:
            pending_scenes.append(scene)
    
    if not pending_scenes:
        print(f"⏩ 所有 {len(cut_scenes)} 个片段已分析")
        return list(existing_analyses.values())
    
    print(f"📦 共 {len(pending_scenes)} 个新片段需要分析")
    
    # 分批：每批最多2个片段，确保每批运行时间 < 10分钟
    BATCH_SIZE = 2
    batches = []
    for i in range(0, len(pending_scenes), BATCH_SIZE):
        batch = pending_scenes[i:i+BATCH_SIZE]
        batches.append(batch)
    
    print(f"🎯 分成 {len(batches)} 批，每批最多 {BATCH_SIZE} 个片段")
    
    # 逐批处理（每批一个独立进程）
    for batch_idx, batch in enumerate(batches, 1):
        print(f"\n{'='*60}")
        print(f"批次 {batch_idx}/{len(batches)}: {len(batch)} 个片段")
        print(f"{'='*60}")
        
        # 准备批次数据
        batch_data = {
            'batch_id': batch_idx,
            'scenes': batch,
            'question': question,
            'output_dir': str(output_dir)
        }
        
        # 保存批次数据到临时文件
        batch_file = output_dir / f".batch_{batch_idx}_input.json"
        with open(batch_file, 'w') as f:
            json.dump(batch_data, f)
        
        # 启动独立进程分析该批次
        import subprocess
        script_path = Path(__file__).parent / "batch_worker.py"
        
        print(f"🚀 启动独立进程分析批次 {batch_idx}...")
        result = subprocess.run(
            [sys.executable, str(script_path), str(batch_file)],
            capture_output=True,
            text=True,
            timeout=600,  # 10分钟超时
            env=os.environ.copy(),
        )
        
        # 输出子进程日志
        if result.stdout:
            print(result.stdout)
        if result.stderr:
            print(f"⚠️ 批次 {batch_idx} 错误: {result.stderr[:200]}")
        
        # 检查批次结果
        batch_result_file = output_dir / f".batch_{batch_idx}_result.json"
        if batch_result_file.exists():
            with open(batch_result_file) as f:
                batch_results = json.load(f)
            
            # 合并结果
            for result in batch_results:
                existing_analyses[result['idx']] = result
            
            # 立即保存总进度
            progress['analyses'] = list(existing_analyses.values())
            save_progress(output_dir, progress)
            
            success_count = len([r for r in batch_results if r['status'] == 'success'])
            print(f"✅ 批次 {batch_idx} 完成: {success_count}/{len(batch_results)} 成功")
            
            # 清理批次临时文件
            batch_file.unlink(missing_ok=True)
            batch_result_file.unlink(missing_ok=True)
        else:
            print(f"❌ 批次 {batch_idx} 未生成结果文件")
    
    print(f"\n{'='*60}")
    print(f"✅ 全部分析完成")
    total_success = len([a for a in existing_analyses.values() if a['status'] == 'success'])
    print(f"   总计: {total_success}/{len(cut_scenes)} 成功")
    print(f"{'='*60}")
    
    return list(existing_analyses.values())


# ============================================================================
# 生成报告
# ============================================================================

def extract_scene_screenshot(scene_path: str, scene_idx: int, screenshots_dir: Path) -> Path | None:
    """为单个镜头抽取一张中间帧截图。"""
    try:
        scene_file = Path(scene_path)
        if not scene_file.exists():
            return None

        screenshots_dir.mkdir(parents=True, exist_ok=True)
        output_path = screenshots_dir / f"scene_{scene_idx:03d}.jpg"
        if output_path.exists() and output_path.stat().st_size > 0:
            return output_path

        probe = subprocess.run(
            [
                "ffprobe", "-v", "error",
                "-show_entries", "format=duration",
                "-of", "default=noprint_wrappers=1:nokey=1",
                str(scene_file),
            ],
            capture_output=True,
            text=True,
            timeout=15,
        )
        duration = 0.0
        if probe.returncode == 0:
            try:
                duration = float((probe.stdout or "0").strip() or 0)
            except Exception:
                duration = 0.0
        seek_time = max(duration / 2, 0.0) if duration > 0 else 0.0

        cmd = [
            "ffmpeg", "-y",
            "-ss", f"{seek_time:.3f}",
            "-i", str(scene_file),
            "-frames:v", "1",
            "-q:v", "2",
            str(output_path),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode == 0 and output_path.exists() and output_path.stat().st_size > 0:
            return output_path
        return None
    except Exception:
        return None


def create_report(video_path: str, scenes: list, analyses: list, output_dir: Path, user_prompt: str = None):
    """生成 Word 报告"""
    print("\n" + "="*60)
    print("步骤4: 生成 Word 报告")
    print("="*60)
    
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # 标题
    title = doc.add_heading('视频分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 统计信息
    doc.add_heading('一、统计信息', 1)
    
    completeness = check_analysis_completeness(analyses)
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    info = [
        ("视频路径", video_path),
        ("镜头总数", str(len(scenes))),
        ("分析片段", str(completeness['total'])),
        ("成功分析", f"{completeness['success']} ({completeness['success_rate']*100:.1f}%)"),
        ("失败", str(completeness['failed'])),
        ("跳过", str(completeness['skipped'])),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    ]
    
    for i, (k, v) in enumerate(info):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    # 完整性说明
    if completeness['success_rate'] < SUCCESS_RATE_THRESHOLD:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("⚠️ 注意：").bold = True
        p.add_run(f"分析成功率 {completeness['success_rate']*100:.1f}% 未达到 {SUCCESS_RATE_THRESHOLD*100}% 阈值，"
                  f"部分片段标记为失败。建议重试失败片段以获得更完整的分析。")
    
    # 分析要求
    if user_prompt:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("分析要求：").bold = True
        p.add_run(user_prompt)
    
    # 逐场景分析
    doc.add_heading('二、逐场景分析', 1)
    screenshots_dir = output_dir / "scene_screenshots"
    
    for item in analyses:
        idx = item['idx']
        scene = item['scene']
        status = item.get('status', 'unknown')
        
        # 标题带状态标记
        status_icon = "✅" if status == "success" else "❌" if status in ['failed', 'error', 'exception'] else "⏭️"
        doc.add_heading(f'{status_icon} 场景 {idx+1}', 2)
        
        # 时间信息
        start = scene.get("start", scene.get("start_time", 0))
        end = scene.get("end", scene.get("end_time", 0))
        size_mb = item.get("size_mb", 0)
        
        p = doc.add_paragraph()
        p.add_run(f"时间: ").bold = True
        p.add_run(f"{start:.1f}s - {end:.1f}s")
        p.add_run(f"  |  大小: ").bold = True
        p.add_run(f"{size_mb:.1f}MB")
        p.add_run(f"  |  状态: ").bold = True
        p.add_run(status)

        # 镜头截图
        screenshot_path = extract_scene_screenshot(item.get('path', ''), idx, screenshots_dir)
        if screenshot_path:
            doc.add_paragraph('镜头截图')
            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_run = pic_p.add_run()
            pic_run.add_picture(str(screenshot_path), width=Inches(4.8))
        
        # 分析内容
        doc.add_heading('分析内容', 3)
        analysis_text = item["analysis"]
        
        if status != "success":
            p = doc.add_paragraph()
            run = p.add_run(analysis_text)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.italic = True
        else:
            doc.add_paragraph(analysis_text)
    
    # 保存
    report_path = output_dir / f"video_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(report_path)
    
    print(f"✅ 报告已保存: {report_path}")
    return report_path


# ============================================================================
# 主函数
# ============================================================================

def main():
    parser = argparse.ArgumentParser(description="视频分析工作流 - V3（95%阈值 + 暂停恢复）")
    parser.add_argument("video", nargs='?', help="输入视频路径")
    parser.add_argument("-o", "--output", default="./output", help="输出目录")
    parser.add_argument("-q", "--question", help="自定义分析问题")
    parser.add_argument("--resume", action="store_true", help="断点续传")
    parser.add_argument("--restart", action="store_true", help="重新开始")
    parser.add_argument("--auto-continue", action="store_true", help="自动继续（不询问）")
    parser.add_argument("--auto-pause", action="store_true", help="自动暂停（不询问）")
    parser.add_argument("--status", action="store_true", help="查看状态")
    
    args = parser.parse_args()
    
    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 查看状态
    if args.status:
        state = load_incomplete_state(output_dir)
        if state:
            print(f"📋 有暂停的分析: {state['video_path']}")
            print(f"   成功率: {state['completeness']['success_rate']*100:.1f}%")
            print(f"   暂停时间: {state['paused_at']}")
        else:
            progress = load_progress(output_dir)
            if progress.get('analyses'):
                print(f"🔄 有进行中的分析: {len(progress['analyses'])} 个片段")
            else:
                print("✅ 没有进行中的分析")
        return 0
    
    # 检查暂停状态
    incomplete_state = load_incomplete_state(output_dir)
    if incomplete_state and not args.restart and not args.video:
        decision = show_resume_prompt(incomplete_state)
        
        if decision == "continue":
            # 直接生成报告
            report_path = create_report(
                incomplete_state['video_path'],
                incomplete_state['scenes'],
                incomplete_state['analyses'],
                output_dir,
                incomplete_state.get('user_prompt')
            )
            clear_incomplete_state(output_dir)
            clear_progress(output_dir)
            print(f"\n✅ 报告已生成: {report_path}")
            return 0
            
        elif decision == "retry":
            # 恢复分析
            args.video = incomplete_state['video_path']
            args.question = incomplete_state.get('user_prompt') or args.question
            args.resume = True
            # 恢复进度
            progress = {
                'video_path': incomplete_state['video_path'],
                'scenes': incomplete_state['scenes'],
                'analyses': incomplete_state['analyses'],
                'scenes_detected': True,
                'scenes_cut': len(incomplete_state['analyses']),
                'completed': False
            }
            save_progress(output_dir, progress)
            clear_incomplete_state(output_dir)
            
        elif decision == "new":
            # 清除状态，开始新分析
            clear_incomplete_state(output_dir)
            clear_progress(output_dir)
            print("🆕 开始新分析")
            if not args.video:
                print("❌ 请提供视频路径")
                return 1
    
    # 需要视频路径
    if not args.video:
        parser.print_help()
        return 1
    
    original_video = resolve_input_video(args.video, output_dir)
    video_id = get_video_id(args.video)
    
    # 设置日志
    global logger
    logger = setup_logging(output_dir, video_id)
    
    # 场景目录
    scenes_dir = output_dir / "scenes"
    scenes_dir.mkdir(exist_ok=True)
    
    # 加载或创建进度
    if args.restart:
        progress = load_progress(output_dir)
        progress['video_path'] = original_video
        progress['completed'] = False
        clear_progress(output_dir)
        clear_incomplete_state(output_dir)
        logger.info("🔄 重新开始（忽略断点）")
    elif args.resume:
        progress = load_progress(output_dir)
        if progress.get('video_path') != original_video:
            logger.warning("⚠️ 断点视频与当前视频不匹配，重新开始")
            progress = {
                'video_path': original_video,
                'preprocessed': False,
                'scenes_detected': False,
                'scenes': [],
                'scenes_cut': 0,
                'analyses': [],
                'completed': False
            }
        else:
            logger.info(f"⏩ 断点续传：已分析 {len(progress.get('analyses', []))} 个片段")
    else:
        progress = {
            'video_path': original_video,
            'preprocessed': False,
            'scenes_detected': False,
            'scenes': [],
            'scenes_cut': 0,
            'analyses': [],
            'completed': False
        }
    
    logger.info("="*60)
    logger.info("视频分析工作流 V3 - 95%阈值 + 暂停恢复")
    logger.info("="*60)
    logger.info(f"📹 原版视频: {original_video}")
    logger.info(f"📁 输出目录: {output_dir}")
    logger.info(f"🎯 成功率阈值: {SUCCESS_RATE_THRESHOLD*100}%")
    logger.info("")
    
    # 步骤0: 预处理
    if progress.get('preprocessed') and not args.restart:
        preprocessed_video = str(scenes_dir / f"preprocessed_{video_id}_480x270.mp4")
        logger.info(f"⏩ 跳过预处理（断点），使用: {preprocessed_video}")
    else:
        preprocessed_video = preprocess_video(original_video, scenes_dir, video_id)
        if preprocessed_video:
            progress['preprocessed'] = True
            save_progress(output_dir, progress)
    
    if not preprocessed_video or not os.path.exists(preprocessed_video):
        logger.error("❌ 预处理失败")
        return 1
    
    # 步骤1: 分镜检测
    if progress.get('scenes_detected') and progress.get('scenes') and not args.restart:
        scenes = progress['scenes']
        logger.info(f"⏩ 跳过分镜检测（断点），使用已检测的 {len(scenes)} 个镜头")
    else:
        scenes = detect_scenes(preprocessed_video, scenes_dir, video_id)
        if scenes:
            progress['scenes_detected'] = True
            progress['scenes'] = scenes
            save_progress(output_dir, progress)
    
    if not scenes:
        logger.error("❌ 分镜检测失败")
        return 1
    
    # 步骤2: 切割
    clips_dir = scenes_dir / f"clips_{video_id}"
    existing_clips = []
    if clips_dir.exists():
        for scene in scenes:
            clip_path = clips_dir / f"scene_{scene['index']:03d}_{scene['start']:.3f}s.mp4"
            if clip_path.exists():
                existing_clips.append({
                    "idx": scene['index'],
                    "path": str(clip_path),
                    "scene": scene,
                    "size_mb": clip_path.stat().st_size / 1024 / 1024
                })
    
    # 如果已有足够片段，跳过切割
    if len(existing_clips) >= len(scenes) and not args.restart:
        cut_results = existing_clips
        logger.info(f"⏩ 跳过切割（断点），使用已切割的 {len(cut_results)} 个片段")
        progress['scenes_cut'] = len(cut_results)
        save_progress(output_dir, progress)
    else:
        cut_results = cut_scenes(original_video, scenes, scenes_dir, video_id)
        if cut_results:
            progress['scenes_cut'] = len(cut_results)
            save_progress(output_dir, progress)
    
    if not cut_results:
        logger.error("❌ 切割失败")
        return 1
    
    # 步骤3: 分析（分批独立进程，避免超时）
    analyses = analyze_all_scenes_batched(cut_results, args.question, progress, output_dir)
    progress['analyses'] = analyses
    save_progress(output_dir, progress)
    
    # 检查完整性
    completeness = check_analysis_completeness(analyses)
    
    print("\n" + "="*60)
    print("📊 分析完整性检查")
    print("="*60)
    print(f"成功率: {completeness['success_rate']*100:.1f}%")
    print(f"阈值: {SUCCESS_RATE_THRESHOLD*100}%")
    print(f"状态: {'✅ 完整' if completeness['complete'] else '❌ 不完整'}")
    print(f"成功: {completeness['success']}/{completeness['total']}")
    print(f"失败: {completeness['failed']}")
    print(f"跳过: {completeness['skipped']}")
    
    # 完整性检查 - 95%阈值
    if not completeness['complete']:
        print("\n⚠️ 未达到 95% 成功率要求")
        
        if args.auto_continue:
            decision = "continue"
            print("[自动继续模式] 生成报告...")
        elif args.auto_pause:
            decision = "pause"
            print("[自动暂停模式] 保存进度...")
        else:
            decision = prompt_user_decision(completeness, output_dir)
        
        if decision == "pause":
            save_incomplete_state(output_dir, original_video, scenes, analyses, completeness, args.question)
            save_progress(output_dir, progress)
            print("\n⏸️ 分析已暂停")
            print("下次运行 --resume 或 --status 查看状态")
            return 0
            
        elif decision == "retry":
            print(f"\n🔄 重试 {len(completeness['can_retry'])} 个失败片段...")
            analyses = retry_failed_scenes(analyses, args.question)
            progress['analyses'] = analyses
            save_progress(output_dir, progress)
            
            # 重新检查
            completeness = check_analysis_completeness(analyses)
            print(f"\n重试后成功率: {completeness['success_rate']*100:.1f}%")
            
            if not completeness['complete'] and not args.auto_continue:
                # 还是不够，再次询问
                decision = prompt_user_decision(completeness, output_dir)
                if decision == "pause":
                    save_incomplete_state(output_dir, original_video, scenes, analyses, completeness, args.question)
                    save_progress(output_dir, progress)
                    return 0
                elif decision == "abort":
                    clear_progress(output_dir)
                    clear_incomplete_state(output_dir)
                    print("\n🗑️ 已放弃分析")
                    return 1
                # continue 或 retry 后还是不够，继续生成报告
        
        elif decision == "abort":
            clear_progress(output_dir)
            clear_incomplete_state(output_dir)
            print("\n🗑️ 已放弃分析，进度已清除")
            return 1
        
        # decision == "continue" 继续生成报告
    
    # 生成报告
    report_path = create_report(original_video, scenes, analyses, output_dir, args.question)
    
    # 清理
    clear_progress(output_dir)
    clear_incomplete_state(output_dir)
    
    print("\n" + "="*60)
    print("✅ 分析流程结束")
    print("="*60)
    print(f"最终成功率: {completeness['success_rate']*100:.1f}%")
    print(f"报告: {report_path}")
    
    return 0


if __name__ == "__main__":
    sys.exit(main())
